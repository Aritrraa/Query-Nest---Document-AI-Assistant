import os
import tempfile
import logging
from pathlib import Path

from dotenv import load_dotenv
from flask import Flask, render_template, request, jsonify, session
from werkzeug.utils import secure_filename

import pandas as pd
import docx
import pypdf

# LLM / vector imports (but we will lazy-load the embedding model)
from groq import Groq
from llama_index.core import VectorStoreIndex, Document

# -------------------------
# Config & Logging
# -------------------------
load_dotenv(".env.local")
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
if not GROQ_API_KEY:
    raise ValueError("Missing GROQ_API_KEY in your .env.local file")

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# -------------------------
# Flask App Setup
# -------------------------
app = Flask(__name__)
app.secret_key = os.urandom(24)

# Upload and size limits
app.config['UPLOAD_FOLDER'] = "uploads"
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100 MB per request
Path(app.config['UPLOAD_FOLDER']).mkdir(parents=True, exist_ok=True)

# Groq client (remote LLM inference)
groq_client = Groq(api_key=GROQ_API_KEY)

# We will lazy-load the HuggingFace embedding model when necessary
_embed_model = None


def get_embed_model():
    """
    Lazy-load and return the HuggingFace embedding wrapper used by LlamaIndex.
    Loading only on-demand avoids heavy RAM usage at process start.
    """
    global _embed_model
    if _embed_model is None:
        logger.info("Loading embedding model (this may use memory)...")
        from llama_index.embeddings.huggingface import HuggingFaceEmbedding
        _embed_model = HuggingFaceEmbedding(model_name="sentence-transformers/all-MiniLM-L6-v2")
        logger.info("Embedding model loaded.")
    return _embed_model


# -------------------------
# File processing utilities
# -------------------------
def process_file(file_path: str) -> str:
    """
    Extract text from supported filetypes: pdf, docx, xls/xlsx.
    Returns extracted text as a single string.
    """
    ext = Path(file_path).suffix.lower()
    try:
        if ext == ".pdf":
            with open(file_path, "rb") as f:
                reader = pypdf.PdfReader(f)
                pages_text = []
                for p in reader.pages:
                    txt = p.extract_text()
                    if txt:
                        pages_text.append(txt)
                return "\n".join(pages_text)
        elif ext == ".docx":
            doc = docx.Document(file_path)
            return "\n".join([p.text for p in doc.paragraphs if p.text])
        elif ext in [".xls", ".xlsx"]:
            df = pd.read_excel(file_path)
            return df.to_string(index=False)
        else:
            raise ValueError(f"Unsupported file extension: {ext}")
    except Exception as e:
        raise RuntimeError(f"Failed to process {file_path}: {str(e)}")


# -------------------------
# Groq wrapper (remote LLM)
# -------------------------
def query_groq(context: str, question: str, max_context_chars: int = 60000) -> str:
    """
    Send a prompt to Groq LLM with a context and question.
    We clip the context to max_context_chars to avoid extremely long prompts.
    """
    # Clip context safely (keep last part of the doc if too long)
    if len(context) > max_context_chars:
        logger.info("Context too long; truncating for LLM prompt.")
        context_to_use = context[-max_context_chars:]
    else:
        context_to_use = context

    prompt = f"Context:\n{context_to_use}\n\nQuestion: {question}\n\nAnswer:"
    completion = groq_client.chat.completions.create(
        messages=[{"role": "user", "content": prompt}],
        model="llama3-70b-8192",
    )
    return completion.choices[0].message.content.strip()


# -------------------------
# Routes
# -------------------------
@app.route("/")
def index():
    return render_template("index.html")


@app.route("/upload", methods=["POST"])
def upload_file():
    """
    Upload handler:
    - Saves uploaded files to uploads/
    - Extracts text and concatenates into a temp text file (not session)
    - Builds a VectorStoreIndex (embedding model is lazy-loaded)
    - Persists index to a temp dir and returns its path
    """
    uploaded_files = request.files.getlist("file")
    if not uploaded_files:
        return jsonify({"error": "No file uploaded"}), 400

    filenames = []
    combined_text_pieces = []

    # Basic per-file validation loop
    for file in uploaded_files:
        filename = secure_filename(file.filename)
        if filename == "":
            continue

        ext = os.path.splitext(filename)[1].lower()
        if ext not in [".pdf", ".docx", ".xlsx", ".xls"]:
            return jsonify({"error": f"Unsupported file type: {ext}"}), 400

        save_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
        file.save(save_path)
        filenames.append(filename)

        try:
            text = process_file(save_path)
            if text:
                combined_text_pieces.append(text)
        except Exception as e:
            logger.exception("Error processing file %s: %s", filename, e)
            return jsonify({"error": str(e)}), 500

    if not combined_text_pieces:
        return jsonify({"error": "No text could be extracted from uploaded files."}), 400

    combined_text = "\n\n".join(combined_text_pieces)

    # Save combined text to a temp file (avoid putting full text in session)
    temp_text_fd, temp_text_path = tempfile.mkstemp(prefix="combined_text_", suffix=".txt")
    os.close(temp_text_fd)
    with open(temp_text_path, "w", encoding="utf-8") as f:
        f.write(combined_text)

    # Persist the text path and filenames in the session (small metadata only)
    session["filenames"] = filenames
    session["text_path"] = temp_text_path

    # Build & persist VectorStoreIndex (embedding model loaded on demand)
    try:
        documents = [Document(text=combined_text)]
        embed_model = get_embed_model()  # lazy-load here
        index = VectorStoreIndex.from_documents(documents, embed_model=embed_model)

        # Persist to a temp directory
        temp_index_dir = tempfile.mkdtemp(prefix="llama_index_")
        index.storage_context.persist(persist_dir=temp_index_dir)

        # store index path in session (only path, index itself not kept in memory)
        session["index_path"] = temp_index_dir
        logger.info("Index persisted to %s", temp_index_dir)
    except Exception as e:
        logger.exception("Failed to create/persist index: %s", e)
        # Clean up text file if index creation failed
        try:
            os.remove(temp_text_path)
        except Exception:
            pass
        return jsonify({"error": f"Failed to create index: {str(e)}"}), 500

    return jsonify({"message": "Files uploaded and processed", "index_path": temp_index_dir})


@app.route("/query", methods=["POST"])
def query_document():
    """
    Query the uploaded documents. We read the combined_text from the temp file
    and send it with the user's question to the Groq LLM.
    """
    data = request.get_json(force=True)
    question = data.get("question", "").strip()
    if not question:
        return jsonify({"error": "Question is required"}), 400

    text_path = session.get("text_path")
    if not text_path or not os.path.exists(text_path):
        return jsonify({"error": "No document uploaded yet"}), 400

    try:
        with open(text_path, "r", encoding="utf-8") as f:
            combined_text = f.read()
    except Exception as e:
        logger.exception("Failed reading combined text: %s", e)
        return jsonify({"error": f"Failed to read stored text: {str(e)}"}), 500

    try:
        answer = query_groq(combined_text, question)
        return jsonify({"response": answer})
    except Exception as e:
        logger.exception("Groq query failed: %s", e)
        return jsonify({"error": str(e)}), 500


@app.route("/summary", methods=["GET"])
def summarize_document():
    """
    Summarize the uploaded documents using the Groq LLM.
    """
    text_path = session.get("text_path")
    if not text_path or not os.path.exists(text_path):
        return jsonify({"error": "No document uploaded yet"}), 400

    try:
        with open(text_path, "r", encoding="utf-8") as f:
            combined_text = f.read()
    except Exception as e:
        logger.exception("Failed reading combined text for summary: %s", e)
        return jsonify({"error": f"Failed to read stored text: {str(e)}"}), 500

    summary_prompt = (
        "Please summarize the following document in a few concise bullet points:\n\n"
        f"{combined_text}"
    )

    try:
        summary = query_groq(combined_text, summary_prompt)
        return jsonify({"summary": summary})
    except Exception as e:
        logger.exception("Groq summary failed: %s", e)
        return jsonify({"error": str(e)}), 500


@app.route("/delete", methods=["POST"])
def delete_files():
    """
    Delete uploaded files, the temp combined text file, and clear the session metadata.
    Does not attempt to delete the persisted index directory (optional).
    """
    filenames = session.get("filenames", [])
    for filename in filenames:
        file_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
            except Exception:
                logger.warning("Failed to remove uploaded file: %s", file_path)

    # Remove combined text file
    text_path = session.get("text_path")
    if text_path and os.path.exists(text_path):
        try:
            os.remove(text_path)
        except Exception:
            logger.warning("Failed to remove temp text file: %s", text_path)

    # Optionally remove persisted index directory
    index_path = session.get("index_path")
    if index_path and os.path.exists(index_path):
        try:
            # remove contents
            for root, dirs, files in os.walk(index_path, topdown=False):
                for name in files:
                    os.remove(os.path.join(root, name))
                for name in dirs:
                    os.rmdir(os.path.join(root, name))
            os.rmdir(index_path)
        except Exception:
            logger.warning("Failed to fully remove index dir: %s", index_path)

    session.clear()
    return jsonify({"message": "All uploaded files and temporary data deleted successfully."})


# -------------------------
# Entrypoint
# -------------------------
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
